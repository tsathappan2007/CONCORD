import io
import re
import logging
from typing import Optional, Dict, Any
from fastapi import HTTPException
import pdfplumber
from docx import Document as DocxDocument
from pydantic import BaseModel, Field
from app.services.llm_orchestrator import LLMOrchestrator

logger = logging.getLogger("concord")

class RelevanceResponse(BaseModel):
    is_contract_related: bool = Field(
        ..., 
        description="True if the text resembles a contract, terms sheet, proposal, service agreement, or negotiation constraint/preference list. False if completely unrelated (e.g. cooking recipe, news article, code, fiction)."
    )
    reason: str = Field(
        ..., 
        description="Brief 1-sentence reason for the classification."
    )

class DocParser:
    def __init__(self, llm_orchestrator: Optional[LLMOrchestrator] = None):
        self.llm_orchestrator = llm_orchestrator or LLMOrchestrator()

    def validate_and_extract(self, file_content: bytes, filename: str, force_ocr: bool = False) -> Dict[str, Any]:
        """
        Full 5-stage validation and extraction pipeline.
        Returns a dictionary with:
          - 'text': Extracted text
          - 'word_count': Word count
          - 'char_count': Character count
          - 'warning': Soft relevance warning (if any)
        """
        # --- Stage 1: Size check (10MB Max) ---
        if len(file_content) > 10 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="File exceeds the 10MB size limit.")

        # --- Stage 2: Format & Signature Validation ---
        ext = self._validate_signature(file_content, filename)

        # --- Stage 3: Extraction / Integrity Check ---
        text = ""
        if force_ocr and ext == "pdf":
            # If the user explicitly opts in for OCR, we use Gemini's multimodal capabilities
            text = self._extract_ocr_gemini(file_content)
        else:
            text = self._extract_raw_text(file_content, ext)

        # --- Stage 4: Readability & Emptiness Validation ---
        text = self._validate_readability(text, ext, force_ocr)

        # --- Stage 5: Content Relevance Validation (Soft Check) ---
        warning = self._check_relevance(text)

        word_count = len(re.findall(r"\w+", text))
        char_count = len(text)

        return {
            "text": text,
            "word_count": word_count,
            "char_count": char_count,
            "warning": warning
        }

    def _validate_signature(self, file_content: bytes, filename: str) -> str:
        ext = filename.split(".")[-1].lower()
        if ext not in ["pdf", "docx", "txt"]:
            raise HTTPException(
                status_code=400, 
                detail="Unsupported file format. Please upload a PDF, DOCX, or TXT file."
            )

        # Check binary headers
        if ext == "pdf":
            if not file_content.startswith(b"%PDF"):
                raise HTTPException(
                    status_code=400, 
                    detail="Invalid file signature. The file claims to be a PDF but does not match PDF signature."
                )
        elif ext == "docx":
            if not file_content.startswith(b"PK\x03\x04"):
                raise HTTPException(
                    status_code=400, 
                    detail="Invalid file signature. The file claims to be a DOCX but does not match DOCX signature."
                )
        
        return ext

    def _extract_raw_text(self, file_content: bytes, ext: str) -> str:
        text = ""
        try:
            if ext == "pdf":
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    for page in pdf.pages:
                        extracted = page.extract_text()
                        if extracted:
                            text += extracted + "\n"
            elif ext == "docx":
                doc = DocxDocument(io.BytesIO(file_content))
                text = "\n".join([p.text for p in doc.paragraphs if p.text])
                # Also extract tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                        if row_text:
                            text += row_text + "\n"
            elif ext == "txt":
                try:
                    text = file_content.decode("utf-8")
                except UnicodeDecodeError:
                    text = file_content.decode("latin-1")
        except Exception as e:
            logger.error(f"Extraction failed: {e}")
            raise HTTPException(
                status_code=400, 
                detail="This file appears to be corrupted or unreadable. Please re-export and re-upload."
            )

        return text

    def _extract_ocr_gemini(self, file_content: bytes) -> str:
        """
        Uses Gemini 2.5 Flash's native multimodal capabilities to transcribe scanned PDF documents.
        This provides high quality OCR with zero external binary dependencies.
        """
        if not self.llm_orchestrator.gemini_client:
            raise HTTPException(
                status_code=400,
                detail="OCR fallback is unavailable because Gemini API is not configured on the backend."
            )
        
        try:
            logger.info("Executing Gemini-based OCR for scanned PDF.")
            # Upload file as inline bytes with pdf MIME type
            from google.genai import types
            
            response = self.llm_orchestrator.gemini_client.models.generate_content(
                model="gemini-2.5-flash",
                contents=[
                    types.Part.from_bytes(
                        data=file_content,
                        mime_type="application/pdf"
                    ),
                    "Read this PDF and perform OCR to extract all the text exactly as written. If there are tables, represent them as Markdown tables."
                ]
            )
            return response.text or ""
        except Exception as e:
            logger.error(f"Gemini OCR extraction failed: {e}")
            raise HTTPException(
                status_code=500,
                detail="Gemini OCR extraction failed. Please re-export the file as text-based or try again."
            )

    def _validate_readability(self, text: str, ext: str, force_ocr: bool) -> str:
        cleaned_text = text.strip()
        word_count = len(re.findall(r"\w+", cleaned_text))
        char_count = len(cleaned_text)

        if char_count == 0 or word_count == 0:
            raise HTTPException(
                status_code=400, 
                detail="This document appears to be blank or contains no contract content. Please upload a complete document."
            )

        # Scanned PDF check
        if ext == "pdf" and word_count < 30 and not force_ocr:
            raise HTTPException(
                status_code=400,
                detail="This PDF appears to contain only scanned images with no readable text. Please upload a text-based document, or check the 'OCR Extract' option."
            )

        # Minimum thresholds
        if word_count < 50 or char_count < 200:
            raise HTTPException(
                status_code=400,
                detail=f"Extracted content is too short ({word_count} words). The document must contain at least 50 words and 200 characters of text."
            )

        return cleaned_text

    def _check_relevance(self, text: str) -> Optional[str]:
        """
        Soft relevance verification via LLM. Returns a warning string if relevance is low.
        """
        sample_text = text[:1500]
        system_prompt = (
            "You are an assistant checking if a document is relevant to a commercial contract, agreement, "
            "terms sheet, service scope, or business negotiation. Return a structured JSON response."
        )
        user_prompt = f"Verify if the following text is contract or negotiation related:\n\n{sample_text}"
        
        try:
            res: RelevanceResponse = self.llm_orchestrator.generate_structured(
                system_prompt, user_prompt, RelevanceResponse
            )
            if not res.is_contract_related:
                return "This document doesn't look like a typical contract. You can continue, but please double-check you uploaded the right file."
        except Exception as e:
            logger.warning(f"Relevance soft check failed to evaluate: {e}")
        
        return None

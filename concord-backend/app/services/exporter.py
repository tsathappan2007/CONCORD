import io
import logging
from datetime import datetime
from typing import List, Dict, Any
from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from app.services.supabase_client import get_supabase_admin

logger = logging.getLogger("concord")

class ContractExporter:
    @staticmethod
    def generate_docx(session_title: str, agreed_terms: List[Dict[str, Any]], logs: List[Dict[str, Any]]) -> bytes:
        """Generates a professional DOCX file of the negotiation log trail."""
        doc = Document()
        
        # Style layout
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Title
        p_title = doc.add_paragraph()
        p_title.alignment = 1 # Center
        r_title = p_title.add_run("Talks of Negotiation")
        r_title.font.name = "Arial"
        r_title.font.size = Pt(22)
        r_title.bold = True
        
        # Meta info
        doc.add_paragraph(f"Session Title: {session_title}")
        doc.add_paragraph(f"Export Date: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}")
        
        p_desc = doc.add_paragraph()
        p_desc.add_run(
            "This document logs the autonomous negotiation steps and justifications "
            "exchanged between the parties' AI agents during the agreement creation process on CONCORD."
        ).italic = True
        
        # Audit Trail
        doc.add_heading("Negotiation Trail & Logs", level=1)
        
        for entry in logs:
            sender = str(entry["sender"]).upper()
            tool = str(entry["tool_called"]).replace("_", " ").title()
            term_suffix = f" on '{entry['term']}'" if entry["term"] else ""
            
            p_log = doc.add_paragraph()
            r_header = p_log.add_run(f"[{sender} - {tool}{term_suffix}]: ")
            r_header.bold = True
            p_log.add_run(str(entry["reasoning"]))
            
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream.getvalue()

    @staticmethod
    def get_party_details(supabase, user_id: str) -> dict:
        """Helper to fetch party email, name and phone from Supabase Auth admin API."""
        if not user_id:
            return {"email": "Party B", "name": "Party B", "phone": "[Client Phone]"}
            
        try:
            user_info = supabase.auth.admin.get_user_by_id(user_id)
            user = user_info.user
            metadata = user.user_metadata or {}
            name = metadata.get("name") or metadata.get("full_name") or user.email or "Party"
            phone = metadata.get("phone") or metadata.get("phone_number") or "[Phone Number]"
            return {
                "email": user.email or "",
                "name": name,
                "phone": phone
            }
        except Exception as e:
            return {"email": "Party", "name": "Party", "phone": "[Phone Number]"}

    @staticmethod
    def generate_pdf(session_title: str, agreed_terms: List[Dict[str, Any]], party_a: dict, party_b: dict) -> bytes:
        """Generates a PDF service agreement using a professional template."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=letter, 
            rightMargin=54, 
            leftMargin=54, 
            topMargin=54, 
            bottomMargin=54
        )
        
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            "DocTitle",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=24,
            leading=28,
            alignment=1, # Center
            spaceAfter=25
        )
        h1_style = ParagraphStyle(
            "DocH1",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=11,
            leading=14,
            spaceBefore=15,
            spaceAfter=8,
            textColor=colors.black
        )
        body_style = ParagraphStyle(
            "DocBody",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=10,
            leading=14,
            spaceAfter=12
        )
        
        story = []
        
        # 1. Main Title
        story.append(Paragraph("SERVICE AGREEMENT", title_style))
        story.append(Spacer(1, 10))
        
        # 2. Heading: PARTIES
        story.append(Paragraph("<u><b>PARTIES</b></u>", h1_style))
        
        effective_date = datetime.utcnow().strftime('%B %d, %Y')
        parties_text = (
            f"This Service Contract Agreement (hereinafter referred to as the <b>\"Agreement\"</b>) is entered into "
            f"on <u>{effective_date}</u> (the <b>\"Effective Date\"</b>), by and between <u>{party_a['name']}</u> "
            f"(Phone: <u>{party_a['phone']}</u>), with an address of <u>[Service Provider Address]</u> (hereinafter referred to as the <b>\"Service Provider\"</b>) "
            f"and <u>{party_b['name']}</u> (Phone: <u>{party_b['phone']}</u>), with an address of <u>[Client Address]</u> (hereinafter referred to as the <b>\"Client\"</b>) "
            f"(collectively referred to as the <b>\"Parties\"</b>)."
        )
        story.append(Paragraph(parties_text, body_style))
        story.append(Spacer(1, 10))
        
        # 3. Heading: LIST OF SERVICES PROVIDED AND THEIR PRICES
        story.append(Paragraph("<u><b>LIST OF SERVICES PROVIDED AND THEIR PRICES</b></u>", h1_style))
        
        services_intro = (
            "During the period of this Agreement, the Service Provider shall have the responsibility to "
            "perform and provide the following services (hereinafter referred to as <b>\"Services\"</b>):"
        )
        story.append(Paragraph(services_intro, body_style))
        
        # List of terms
        term_idx = 1
        payment_days = "30"
        for item in agreed_terms:
            t_name = str(item["term"]).replace("_", " ").title()
            val = str(item["value"])
            
            # Check if this term specifies invoice payment days
            if "payment" in item["term"].lower() or "days" in item["term"].lower() or "invoice" in item["term"].lower():
                if val.isdigit():
                    payment_days = val
            
            story.append(Paragraph(f"{term_idx}. {t_name} (Price/Value: {val})", body_style))
            term_idx += 1
            
        story.append(Spacer(1, 10))
        story.append(Paragraph("The Services are to be paid for as follows:", body_style))
        story.append(Paragraph("Amount at signing of this Agreement: [As negotiated / See terms above]", body_style))
        story.append(Paragraph("Amount at the completion of the provision of the Services: [As negotiated / See terms above]", body_style))
        story.append(Spacer(1, 10))
        
        # 4. Heading: INVOICES
        story.append(Paragraph("<u><b>INVOICES</b></u>", h1_style))
        invoices_text = (
            f"The Parties agree that the invoiced amounts must be paid within <u>{payment_days}</u> days after "
            f"the Client receives the invoice."
        )
        story.append(Paragraph(invoices_text, body_style))
        story.append(Spacer(1, 10))
        
        # 5. Heading: TERM OF AGREEMENT
        story.append(Paragraph("<u><b>TERM OF AGREEMENT</b></u>", h1_style))
        term_text = (
            "This Agreement shall commence on the Effective Date and shall remain in effect until the completion of "
            "the Services or as otherwise terminated in accordance with the terms of this Agreement."
        )
        story.append(Paragraph(term_text, body_style))
        story.append(Spacer(1, 15))
        
        # 6. Signature panels
        story.append(Paragraph("<b>IN WITNESS WHEREOF</b>, the Parties hereto have signed and executed this Agreement.", body_style))
        story.append(Spacer(1, 10))
        
        sig_data = [
            [
                Paragraph(f"<b>Service Provider:</b><br/>{party_a['name']}<br/><br/>_______________________<br/>Authorized Signatory", body_style),
                Paragraph(f"<b>Client:</b><br/>{party_b['name']}<br/><br/>_______________________<br/>Authorized Signatory", body_style)
            ]
        ]
        sig_table = Table(sig_data, colWidths=[240, 240])
        sig_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ]))
        story.append(sig_table)
        
        # Build document with page border decorator
        def draw_agreement_border(canvas, doc_obj):
            canvas.saveState()
            canvas.setStrokeColor(colors.HexColor("#0ea5e9")) # Sky-500
            canvas.setLineWidth(14) # Thick blue border
            canvas.rect(20, 20, doc_obj.pagesize[0] - 40, doc_obj.pagesize[1] - 40)
            canvas.restoreState()
            
        doc.build(story, onFirstPage=draw_agreement_border, onLaterPages=draw_agreement_border)
        buffer.seek(0)
        return buffer.getvalue()

    @classmethod
    def export_and_upload(cls, session_id: str, format_type: str = "pdf") -> str:
        """
        Generates the contract, uploads it to Supabase Storage, 
        and returns the public/signed URL of the file.
        """
        supabase = get_supabase_admin()
        
        # Fetch Session title
        session = supabase.table("negotiation_sessions").select("title").eq("id", session_id).execute().data[0]
        title = session["title"]
        
        # Fetch agreed terms
        terms = supabase.table("agreed_terms").select("*").eq("session_id", session_id).order("term").execute().data
        
        # Fetch logs
        logs = supabase.table("negotiation_logs").select("*").eq("session_id", session_id).order("created_at").execute().data

        # Fetch profiles of Party A and Party B
        try:
            sess_data = supabase.table("negotiation_sessions").select("party_a_id, party_b_id").eq("id", session_id).execute().data[0]
        except Exception as e:
            logger.warning(f"Error fetching party ids for exporter: {e}")
            sess_data = {}

        party_a = cls.get_party_details(supabase, sess_data.get("party_a_id"))
        party_b = cls.get_party_details(supabase, sess_data.get("party_b_id"))

        if format_type == "docx":
            file_bytes = cls.generate_docx(title, terms, logs)
            file_name = "Talks of Negotiation.docx"
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            file_bytes = cls.generate_pdf(title, terms, party_a, party_b)
            file_name = f"concord_agreement_{session_id}.pdf"
            mime_type = "application/pdf"

        # Upload to Supabase Storage bucket 'contracts'
        # Check bucket existence first or just try to upload
        try:
            buckets = supabase.storage.list_buckets()
            if not any(b.name == "contracts" for b in buckets):
                supabase.storage.create_bucket("contracts", options={"public": True})
        except Exception as e:
            logger.warning(f"Error handling bucket check: {e}")

        # Upload
        path_in_bucket = f"{session_id}/{file_name}"
        try:
            supabase.storage.from_("contracts").upload(
                path=path_in_bucket,
                file=file_bytes,
                file_options={"content-type": mime_type, "x-upsert": "true"}
            )
        except Exception as e:
            logger.debug(f"Bucket upload notice: {e}")

        # Get public url
        url = supabase.storage.from_("contracts").get_public_url(path_in_bucket)
        return url
